import docx
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn

import string

wdFormatPDF = 17
def docu(final_summa,agenda):

    doc = docx.Document()
    sec_pr = doc.sections[0]._sectPr # get the section properties el
    # create new borders el
    pg_borders = OxmlElement('w:pgBorders')
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '10') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '10')
        border_el.set(qn('w:color'), 'black')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section

    #agenda=input("Enter Agenda")
    doc.add_heading('Agenda -' + agenda , 0)

    p = doc.add_paragraph(final_summa)
    p.add_run(' :- ').bold = True



    path_doc = string.capwords(agenda)
    path_doc = path_doc.replace(" ", "") + ".docx"
    print(path_doc)

    doc.save(path_doc)


